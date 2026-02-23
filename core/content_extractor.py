"""
Document Content Extraction Module

Extracts text content from various document formats (.txt, .docx, .pdf)
for classification processing. Does not store content - only returns
extracted text as string for immediate processing by classifier.

Supports:
  - .txt files (plain text)
  - .docx files (Microsoft Word documents)
  - .pdf files (Adobe PDF)

Enforces:
  - 5MB maximum file size
  - Skips unsupported file types
  - Graceful error handling
"""

import os
from pathlib import Path
from typing import Optional, Tuple


def extract_text(file_path: str, max_size_mb: int = 5) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text content from a document file.
    
    Args:
        file_path (str): Absolute path to the document file
        max_size_mb (int): Maximum file size in MB (default: 5)
    
    Returns:
        Tuple[Optional[str], Optional[str]]: 
            - (extracted_text, None) on success
            - (None, error_message) on failure
    
    Supports: .txt, .docx, .pdf
    
    Example:
        text, error = extract_text("/path/to/document.pdf")
        if error:
            print(f"Extraction failed: {error}")
        else:
            classification_result = classifier.classify(text)
    """
    
    # Validate file path exists
    if not os.path.exists(file_path):
        return None, f"File not found: {file_path}"
    
    # Get file extension
    file_ext = Path(file_path).suffix.lower()
    
    # Validate file type
    supported_types = {'.txt', '.docx', '.pdf'}
    if file_ext not in supported_types:
        return None, f"Unsupported file type: {file_ext}. Supported: {supported_types}"
    
    # Check file size
    try:
        file_size_bytes = os.path.getsize(file_path)
        file_size_mb = file_size_bytes / (1024 * 1024)
        
        if file_size_mb > max_size_mb:
            return None, f"File too large: {file_size_mb:.2f}MB exceeds {max_size_mb}MB limit"
    except OSError as e:
        return None, f"Cannot access file: {str(e)}"
    
    # Route to appropriate extractor
    try:
        if file_ext == '.txt':
            text = _extract_txt(file_path)
        elif file_ext == '.docx':
            text = _extract_docx(file_path)
        elif file_ext == '.pdf':
            text = _extract_pdf(file_path)
        
        if not text or not text.strip():
            return None, f"No text content extracted from {file_ext} file"
        
        return text, None
    
    except Exception as e:
        return None, f"Extraction error ({file_ext}): {str(e)}"


def _extract_txt(file_path: str) -> str:
    """
    Extract text from .txt file.
    
    Args:
        file_path (str): Path to .txt file
    
    Returns:
        str: Extracted text content
    
    Raises:
        Exception: If file cannot be read
    """
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def _extract_docx(file_path: str) -> str:
    """
    Extract text from .docx file (Microsoft Word).
    
    Args:
        file_path (str): Path to .docx file
    
    Returns:
        str: Extracted text content
    
    Raises:
        Exception: If file is corrupted or cannot be parsed
    """
    from docx import Document
    
    doc = Document(file_path)
    
    # Extract paragraphs
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)


def _extract_pdf(file_path: str) -> str:
    """
    Extract text from .pdf file (Adobe PDF).
    
    Args:
        file_path (str): Path to .pdf file
    
    Returns:
        str: Extracted text content
    
    Raises:
        Exception: If PDF is encrypted or corrupted
    """
    from PyPDF2 import PdfReader
    
    pdf_reader = PdfReader(file_path)
    
    # Check if PDF is encrypted
    if pdf_reader.is_encrypted:
        # Try with empty password
        if not pdf_reader.decrypt(''):
            raise Exception("PDF is password-protected and cannot be extracted")
    
    # Extract text from all pages
    text_parts = []
    for page_num, page in enumerate(pdf_reader.pages):
        try:
            page_text = page.extract_text()
            if page_text.strip():
                text_parts.append(page_text)
        except Exception as e:
            # Log page error but continue
            pass
    
    return "\n".join(text_parts)


def get_supported_formats() -> set:
    """
    Get set of supported file formats.
    
    Returns:
        set: Supported file extensions (e.g., {'.txt', '.docx', '.pdf'})
    """
    return {'.txt', '.docx', '.pdf'}


def is_supported_format(file_path: str) -> bool:
    """
    Check if file format is supported.
    
    Args:
        file_path (str): Path to file
    
    Returns:
        bool: True if format is supported, False otherwise
    """
    file_ext = Path(file_path).suffix.lower()
    return file_ext in get_supported_formats()


if __name__ == "__main__":
    # Simple CLI for testing
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python content_extractor.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    print(f"\nExtracting text from: {file_path}")
    print("-" * 80)
    
    text, error = extract_text(file_path)
    
    if error:
        print(f"❌ Error: {error}")
        sys.exit(1)
    
    print(f"✅ Extraction successful ({len(text)} characters)")
    print("\nExtracted Text (first 500 characters):")
    print("-" * 80)
    print(text[:500])
    if len(text) > 500:
        print(f"\n... ({len(text) - 500} more characters)")
